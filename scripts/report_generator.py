# report_generator.py
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_word_report():
    # Crear nuevo documento
    doc = Document()
    
    # Configurar estilos del documento
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Añadir título
    title = doc.add_heading('Informe de Resúmenes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir fecha
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Procesar cada archivo de resumen
    summaries_dir = Path('summaries')
    if not summaries_dir.exists():
        raise Exception("No se encuentra el directorio de resúmenes")
    
    for summary_file in sorted(summaries_dir.glob('*.txt')):
        # Añadir nombre del archivo como subtítulo
        doc.add_heading(f'Resumen: {summary_file.stem}', level=1)
        
        # Leer y añadir contenido del resumen
        with open(summary_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
            # Separar el contenido por secciones
            sections = content.split('\n\n')
            
            for section in sections:
                if section.strip():
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(section.strip())
        
        # Añadir separador
        doc.add_paragraph('_' * 50)
    
    # Crear directorio para reportes si no existe
    os.makedirs('reports', exist_ok=True)
    
    # Guardar documento
    report_path = f'reports/Informe_Resúmenes_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(report_path)
    print(f"Informe generado: {report_path}")

if __name__ == "__main__":
    create_word_report()